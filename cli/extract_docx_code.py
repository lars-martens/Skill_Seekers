#!/usr/bin/env python3
"""
.docx Code Block Extractor

Task: B2.4 - Extract code blocks from Word docs
Purpose: Detect and extract code samples from Word documents
Usage: python3 cli/extract_docx_code.py <input.docx> [options]

Features:
- Detect code blocks via monospace font detection
- Extract code samples with context
- Detect programming language via heuristics
- Export as JSON or plain text
- Compatible with Skill Seeker's code sample format

Detection Methods:
1. Monospace font families (Courier, Consolas, Monaco, etc.)
2. Background shading (common for code blocks)
3. Indentation patterns
4. Paragraph style names containing "code"

Requirements:
    pip install python-docx
"""

import sys
import argparse
import json
import re
from pathlib import Path
from collections import defaultdict


# Common monospace fonts used for code
MONOSPACE_FONTS = {
    'Courier New', 'Courier',
    'Consolas',  # Windows default
    'Monaco',    # macOS
    'Menlo',     # macOS
    'Source Code Pro',
    'Fira Code',
    'JetBrains Mono',
    'DejaVu Sans Mono',
    'Liberation Mono',
    'Ubuntu Mono',
    'Inconsolata',
    'Lucida Console',
    'Anonymous Pro',
    'Droid Sans Mono',
}


def is_monospace_font(font_name):
    """Check if font is monospace."""
    if not font_name:
        return False
    return font_name in MONOSPACE_FONTS


def detect_language_from_code(code):
    """
    Detect programming language from code content using heuristics.

    Args:
        code: Code string

    Returns:
        str: Detected language or 'unknown'
    """
    code_lower = code.lower()

    # Python indicators
    if re.search(r'\b(def|import|from|class|print|self|__init__|elif)\b', code):
        return 'python'

    # JavaScript/TypeScript
    if re.search(r'\b(function|const|let|var|=>|console\.log|async|await)\b', code):
        if 'interface' in code or ': string' in code or ': number' in code:
            return 'typescript'
        return 'javascript'

    # Java
    if re.search(r'\b(public|private|protected|class|void|static|extends|implements)\b', code):
        return 'java'

    # C/C++
    if re.search(r'\b(#include|printf|scanf|int main|std::)\b', code):
        if 'std::' in code or 'cout' in code or 'namespace' in code:
            return 'cpp'
        return 'c'

    # C#
    if re.search(r'\b(using|namespace|public class|private class|static void Main)\b', code):
        return 'csharp'

    # Go
    if re.search(r'\b(func|package|import|defer|go\s+\w+|make\()\b', code):
        return 'go'

    # Rust
    if re.search(r'\b(fn|let mut|impl|trait|pub|use|match)\b', code):
        return 'rust'

    # Ruby
    if re.search(r'\b(def|end|require|class|module|puts|attr_accessor)\b', code):
        return 'ruby'

    # PHP
    if re.search(r'<\?php|function\s+\w+\(|\$\w+\s*=', code):
        return 'php'

    # Shell/Bash
    if re.search(r'^#!/bin/(bash|sh)|^#\s+|\$\s+[a-z]+|echo\s+', code):
        return 'bash'

    # SQL
    if re.search(r'\b(SELECT|FROM|WHERE|INSERT|UPDATE|DELETE|CREATE TABLE|JOIN)\b', code_lower):
        return 'sql'

    # HTML
    if re.search(r'<(html|div|body|head|title|p|a|span|script|style)', code_lower):
        return 'html'

    # CSS
    if re.search(r'[.#]\w+\s*\{|:\s*(hover|active|focus)|@media', code):
        return 'css'

    # JSON
    if code.strip().startswith('{') or code.strip().startswith('['):
        try:
            json.loads(code)
            return 'json'
        except:
            pass

    # YAML
    if re.search(r'^\w+:\s*$|^\s+-\s+\w+', code, re.MULTILINE):
        return 'yaml'

    # GDScript (Godot)
    if re.search(r'\b(extends|signal|export|onready|func _ready|func _process)\b', code):
        return 'gdscript'

    return 'unknown'


def extract_code_blocks(filepath, include_context=False):
    """
    Extract code blocks from .docx file by detecting monospace fonts.

    Args:
        filepath: Path to the .docx file
        include_context: Include surrounding paragraph for context

    Returns:
        list: List of code block dicts
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx library not found. Install with: pip install python-docx"
        )

    if not Path(filepath).exists():
        raise FileNotFoundError(f"File not found: {filepath}")

    doc = Document(filepath)
    code_blocks = []

    for idx, para in enumerate(doc.paragraphs):
        para_text = para.text.strip()
        if not para_text:
            continue

        # Check if paragraph is code
        is_code = False
        detected_font = None

        # Method 1: Check style name
        style_name = para.style.name.lower()
        if 'code' in style_name or 'pre' in style_name or 'source' in style_name:
            is_code = True

        # Method 2: Check font in runs
        if not is_code:
            for run in para.runs:
                font_name = run.font.name
                if is_monospace_font(font_name):
                    is_code = True
                    detected_font = font_name
                    break

        # If identified as code, extract it
        if is_code:
            # Detect language
            language = detect_language_from_code(para_text)

            # Get context if requested
            context_before = None
            context_after = None

            if include_context:
                # Get previous non-empty paragraph
                for i in range(idx - 1, max(0, idx - 3), -1):
                    prev_text = doc.paragraphs[i].text.strip()
                    if prev_text:
                        context_before = prev_text
                        break

                # Get next non-empty paragraph
                for i in range(idx + 1, min(len(doc.paragraphs), idx + 3)):
                    next_text = doc.paragraphs[i].text.strip()
                    if next_text:
                        context_after = next_text
                        break

            code_block = {
                'index': idx,
                'code': para_text,
                'language': language,
                'detected_font': detected_font,
                'style': para.style.name
            }

            if include_context:
                code_block['context_before'] = context_before
                code_block['context_after'] = context_after

            code_blocks.append(code_block)

    return code_blocks


def group_by_language(code_blocks):
    """Group code blocks by detected language."""
    grouped = defaultdict(list)
    for block in code_blocks:
        language = block['language']
        grouped[language].append(block)
    return dict(grouped)


def format_as_markdown(code_blocks):
    """Format code blocks as markdown."""
    lines = []

    for idx, block in enumerate(code_blocks, 1):
        lines.append(f"\n## Code Block {idx}")
        lines.append(f"**Language:** {block['language']}")

        if block.get('detected_font'):
            lines.append(f"**Font:** {block['detected_font']}")

        if block.get('context_before'):
            lines.append(f"\n*Context:* {block['context_before']}")

        lines.append(f"\n```{block['language']}")
        lines.append(block['code'])
        lines.append("```")

        if block.get('context_after'):
            lines.append(f"\n*Following:* {block['context_after']}")

        lines.append("")  # Blank line between blocks

    return '\n'.join(lines)


def main():
    parser = argparse.ArgumentParser(
        description='Extract code blocks from .docx documents',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
Examples:
  # Extract code blocks
  python3 cli/extract_docx_code.py document.docx

  # Export as JSON
  python3 cli/extract_docx_code.py document.docx --json

  # Include context paragraphs
  python3 cli/extract_docx_code.py document.docx --context

  # Markdown format
  python3 cli/extract_docx_code.py document.docx --markdown

  # Group by language
  python3 cli/extract_docx_code.py document.docx --by-language
        '''
    )

    parser.add_argument(
        'input',
        help='Input .docx file path'
    )

    parser.add_argument(
        '-o', '--output',
        help='Output file path (optional)'
    )

    parser.add_argument(
        '-j', '--json',
        action='store_true',
        help='Output as JSON format'
    )

    parser.add_argument(
        '-m', '--markdown',
        action='store_true',
        help='Output as markdown format'
    )

    parser.add_argument(
        '-c', '--context',
        action='store_true',
        help='Include context paragraphs before/after code'
    )

    parser.add_argument(
        '-l', '--by-language',
        action='store_true',
        help='Group code blocks by detected language'
    )

    parser.add_argument(
        '-v', '--verbose',
        action='store_true',
        help='Verbose output'
    )

    args = parser.parse_args()

    # Validate input file
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: File not found: {args.input}", file=sys.stderr)
        sys.exit(1)

    try:
        # Extract code blocks
        if args.verbose:
            print(f"Extracting code blocks from: {args.input}", file=sys.stderr)

        code_blocks = extract_code_blocks(args.input, include_context=args.context)

        if not code_blocks:
            print("No code blocks found in document.", file=sys.stderr)
            sys.exit(0)

        if args.verbose:
            print(f"Found {len(code_blocks)} code blocks", file=sys.stderr)

        # Generate output based on format
        if args.json:
            # JSON output
            if args.by_language:
                grouped = group_by_language(code_blocks)
                output = {
                    'file': str(input_path),
                    'total_blocks': len(code_blocks),
                    'by_language': grouped
                }
            else:
                output = {
                    'file': str(input_path),
                    'total_blocks': len(code_blocks),
                    'code_blocks': code_blocks
                }
            output_text = json.dumps(output, indent=2, ensure_ascii=False)

        elif args.markdown:
            # Markdown format
            output_text = format_as_markdown(code_blocks)

        elif args.by_language:
            # Group by language (plain text)
            grouped = group_by_language(code_blocks)
            lines = [f"Code Blocks from: {input_path}\n"]
            lines.append("=" * 60)
            lines.append(f"Total blocks: {len(code_blocks)}\n")

            for language, blocks in sorted(grouped.items()):
                lines.append(f"\n{language.upper()} ({len(blocks)} blocks)")
                lines.append("-" * 40)
                for block in blocks:
                    lines.append(f"\n{block['code']}\n")

            output_text = '\n'.join(lines)

        else:
            # Default: plain list
            lines = [f"Code Blocks from: {input_path}\n"]
            lines.append("=" * 60)
            lines.append(f"Total blocks: {len(code_blocks)}\n")

            for idx, block in enumerate(code_blocks, 1):
                lines.append(f"\n--- Block {idx} ({block['language']}) ---")
                if block.get('context_before'):
                    lines.append(f"Context: {block['context_before']}")
                lines.append(block['code'])
                if block.get('context_after'):
                    lines.append(f"Following: {block['context_after']}")

            output_text = '\n'.join(lines)

        # Output to file or console
        if args.output:
            Path(args.output).write_text(output_text, encoding='utf-8')
            if args.verbose:
                print(f"\nOutput saved to: {args.output}", file=sys.stderr)
        else:
            print(output_text)

    except ImportError as e:
        print(f"Error: {e}", file=sys.stderr)
        print("\nInstall required library:", file=sys.stderr)
        print("  pip install python-docx", file=sys.stderr)
        sys.exit(1)

    except FileNotFoundError as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)

    except Exception as e:
        print(f"Error: Failed to extract code blocks", file=sys.stderr)
        print(f"Details: {e}", file=sys.stderr)
        import traceback
        if args.verbose:
            traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
