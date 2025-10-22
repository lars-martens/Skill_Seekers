#!/usr/bin/env python3
"""
Word Document to Claude Skill Converter

Task: B2.6 - Create complete .docx scraper CLI tool
Purpose: Convert Word documents (.docx) into Claude AI skills

Usage:
    python3 cli/docx_scraper.py document.docx --name myskill
    python3 cli/docx_scraper.py document.docx --config configs/doc_config.json
    python3 cli/docx_scraper.py --interactive

Features:
- Extract text, headings, code blocks, and tables from .docx
- Auto-categorize content based on heading structure
- Generate SKILL.md and reference files
- Compatible with Skill Seeker's build pipeline
- Similar output format to doc_scraper.py

Requirements:
    pip install python-docx
"""

import os
import sys
import json
import argparse
import re
import hashlib
from pathlib import Path
from collections import defaultdict
from datetime import datetime


# Import extraction functions from our B2.2-B2.5 tools
# For now, inline the critical functions


def is_monospace_font(font_name):
    """Check if font is monospace (for code detection)."""
    if not font_name:
        return False
    monospace_fonts = {
        'Courier New', 'Courier', 'Consolas', 'Monaco', 'Menlo',
        'Source Code Pro', 'Fira Code', 'JetBrains Mono',
        'DejaVu Sans Mono', 'Liberation Mono', 'Ubuntu Mono',
        'Inconsolata', 'Lucida Console'
    }
    return font_name in monospace_fonts


def detect_language_from_code(code):
    """Detect programming language from code content."""
    code_lower = code.lower()

    # Language patterns (simplified from B2.4)
    if re.search(r'\b(def|import|from|class|print|self|__init__)\b', code):
        return 'python'
    if re.search(r'\b(function|const|let|var|=>|console\.log)\b', code):
        return 'javascript'
    if re.search(r'\b(public|private|class|void|static)\b', code):
        return 'java'
    if re.search(r'<(html|div|body|head|title)', code_lower):
        return 'html'
    if re.search(r'\b(SELECT|FROM|WHERE|INSERT)\b', code_lower):
        return 'sql'

    return 'unknown'


class DocxToSkillConverter:
    """Convert .docx documents to Claude skills."""

    def __init__(self, config):
        """
        Initialize converter with configuration.

        Args:
            config: Dict with keys:
                - name: Skill name
                - description: Skill description
                - files: List of .docx file paths
                - categories: Optional category keywords
        """
        self.config = config
        self.name = config['name']
        self.files = config.get('files', [])
        self.description = config.get('description', '')

        # Paths
        self.data_dir = Path(f"output/{self.name}_data")
        self.skill_dir = Path(f"output/{self.name}")

        # State
        self.pages = []
        self.categories = defaultdict(list)

        # Create directories
        self.data_dir.mkdir(parents=True, exist_ok=True)
        (self.data_dir / "pages").mkdir(exist_ok=True)
        self.skill_dir.mkdir(parents=True, exist_ok=True)
        (self.skill_dir / "references").mkdir(exist_ok=True)
        (self.skill_dir / "scripts").mkdir(exist_ok=True)
        (self.skill_dir / "assets").mkdir(exist_ok=True)

    def extract_from_docx(self, filepath):
        """
        Extract all content from a .docx file.

        Returns:
            dict: Page data with text, headings, code, tables
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx library not found. Install with: pip install python-docx"
            )

        if not Path(filepath).exists():
            raise FileNotFoundError(f"File not found: {filepath}")

        print(f"  📄 Processing: {Path(filepath).name}")

        doc = Document(filepath)

        # Extract document title (first Heading 1 or filename)
        title = Path(filepath).stem
        for para in doc.paragraphs:
            if para.style.name == 'Heading 1' and para.text.strip():
                title = para.text.strip()
                break

        # Extract all content
        content_parts = []
        headings = []
        code_blocks = []
        tables_data = []

        current_section = None

        for idx, para in enumerate(doc.paragraphs):
            para_text = para.text.strip()
            if not para_text:
                continue

            style_name = para.style.name

            # Track headings
            if style_name.startswith('Heading'):
                try:
                    level = int(style_name.split()[-1])
                except (ValueError, IndexError):
                    level = 1

                headings.append({
                    'level': level,
                    'text': para_text,
                    'index': idx
                })

                if level == 1:
                    current_section = para_text.lower().replace(' ', '_')

                content_parts.append(f"\n{'#' * level} {para_text}\n")
                continue

            # Check if paragraph is code
            is_code = False
            style_lower = style_name.lower()

            if 'code' in style_lower or 'pre' in style_lower:
                is_code = True
            else:
                # Check font
                for run in para.runs:
                    if is_monospace_font(run.font.name):
                        is_code = True
                        break

            if is_code:
                language = detect_language_from_code(para_text)
                code_blocks.append({
                    'language': language,
                    'code': para_text
                })
                content_parts.append(f"\n```{language}\n{para_text}\n```\n")
            else:
                content_parts.append(para_text)

        # Extract tables
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)

            if rows:
                tables_data.append({
                    'rows': rows,
                    'row_count': len(rows),
                    'column_count': len(rows[0]) if rows else 0
                })

                # Add markdown table to content
                content_parts.append("\n")
                if len(rows) > 0:
                    # Header
                    content_parts.append("| " + " | ".join(rows[0]) + " |")
                    content_parts.append("| " + " | ".join(["---"] * len(rows[0])) + " |")
                    # Data rows
                    for row in rows[1:]:
                        content_parts.append("| " + " | ".join(row) + " |")
                content_parts.append("\n")

        # Build full content
        content = '\n\n'.join(content_parts)

        # Determine category from headings
        category = self._categorize_from_headings(headings)

        return {
            'url': f'file:///{filepath}',
            'title': title,
            'content': content,
            'headings': headings,
            'code_samples': code_blocks,
            'tables': tables_data,
            'patterns': [],  # Could extract patterns in future
            'category': category,
            'source_file': str(filepath)
        }

    def _categorize_from_headings(self, headings):
        """Determine category from heading structure."""
        if not headings:
            return 'general'

        # Use first H1 as category
        for heading in headings:
            if heading['level'] == 1:
                return heading['text'].lower().replace(' ', '_')

        return 'general'

    def process_all_files(self):
        """Process all .docx files and extract content."""
        print(f"\n🔍 Processing {len(self.files)} file(s)...")

        for filepath in self.files:
            try:
                page_data = self.extract_from_docx(filepath)
                self.pages.append(page_data)

                # Save individual page
                page_id = hashlib.md5(str(filepath).encode()).hexdigest()[:12]
                page_file = self.data_dir / "pages" / f"{page_id}.json"

                with open(page_file, 'w', encoding='utf-8') as f:
                    json.dump(page_data, f, indent=2, ensure_ascii=False)

                print(f"    ✓ {Path(filepath).name} ({len(page_data['code_samples'])} code blocks, {len(page_data['tables'])} tables)")

            except Exception as e:
                print(f"    ✗ Failed to process {filepath}: {e}")
                continue

        # Save summary
        summary = {
            'name': self.name,
            'description': self.description,
            'files_processed': len(self.pages),
            'total_code_blocks': sum(len(p['code_samples']) for p in self.pages),
            'total_tables': sum(len(p['tables']) for p in self.pages),
            'timestamp': datetime.now().isoformat()
        }

        with open(self.data_dir / "summary.json", 'w', encoding='utf-8') as f:
            json.dump(summary, f, indent=2)

        print(f"\n✅ Extracted {len(self.pages)} document(s)")

    def build_skill(self):
        """Build skill from extracted data."""
        print(f"\n🔨 Building skill: {self.name}")

        if not self.pages:
            print("❌ No pages to build from. Run extraction first.")
            return

        # Group pages by category
        categories = defaultdict(list)
        for page in self.pages:
            category = page['category']
            categories[category].append(page)

        # Create references for each category
        for category, pages in categories.items():
            ref_file = self.skill_dir / "references" / f"{category}.md"

            lines = [f"# {category.replace('_', ' ').title()}\n"]

            for page in pages:
                lines.append(f"## {page['title']}\n")
                lines.append(f"Source: `{Path(page['source_file']).name}`\n")
                lines.append(page['content'])
                lines.append("\n---\n")

            ref_file.write_text('\n'.join(lines), encoding='utf-8')
            print(f"  📝 Created: references/{category}.md")

        # Create references index
        index_lines = [f"# {self.name.title()} Documentation\n"]
        index_lines.append(f"{self.description}\n")
        index_lines.append("## Categories\n")

        for category in sorted(categories.keys()):
            page_count = len(categories[category])
            index_lines.append(f"- [{category.replace('_', ' ').title()}]({category}.md) ({page_count} document(s))")

        (self.skill_dir / "references" / "index.md").write_text(
            '\n'.join(index_lines),
            encoding='utf-8'
        )

        # Create SKILL.md
        self._create_skill_md(categories)

        print(f"\n✅ Skill built successfully: {self.skill_dir}/")
        print(f"\nNext steps:")
        print(f"  1. Review: cat {self.skill_dir}/SKILL.md")
        print(f"  2. Package: python3 cli/package_skill.py {self.skill_dir}/")

    def _create_skill_md(self, categories):
        """Create SKILL.md file."""
        # Collect code examples
        all_code = []
        for page in self.pages:
            all_code.extend(page['code_samples'])

        # Get best examples (first 5)
        code_examples = all_code[:5]

        skill_content = f'''# {self.name.title()} Skill

{self.description}

## When to Use This Skill

Use this skill when working with {self.name} documentation, code examples, or implementation questions.

## What This Skill Knows

This skill contains:
- Complete documentation from {len(self.pages)} document(s)
- {len(all_code)} code examples
- {sum(len(p['tables']) for p in self.pages)} reference tables
- {len(categories)} main categories

## Quick Reference

### Code Examples

'''

        # Add code examples
        for idx, code in enumerate(code_examples, 1):
            skill_content += f"**Example {idx}** ({code['language']}):\n"
            skill_content += f"```{code['language']}\n"
            skill_content += code['code']
            skill_content += "\n```\n\n"

        skill_content += "\n## Categories\n\n"

        for category in sorted(categories.keys()):
            skill_content += f"- **{category.replace('_', ' ').title()}**: See `references/{category}.md`\n"

        skill_content += f'''

## How to Navigate

1. **Start with `references/index.md`** - Overview of all categories
2. **Explore categories** - Each category has its own reference file
3. **Search code examples** - Look for specific patterns in references

## Reference Files

'''

        for category in sorted(categories.keys()):
            skill_content += f"- `references/{category}.md`\n"

        (self.skill_dir / "SKILL.md").write_text(skill_content, encoding='utf-8')
        print(f"  📄 Created: SKILL.md")


def load_config_file(config_path):
    """Load configuration from JSON file."""
    with open(config_path, 'r', encoding='utf-8') as f:
        return json.load(f)


def interactive_mode():
    """Interactive configuration wizard."""
    print("📝 Word Document to Skill Converter\n")

    name = input("Skill name: ").strip()
    description = input("Description: ").strip()

    files = []
    print("\nEnter .docx file paths (one per line, empty line to finish):")
    while True:
        filepath = input("  File: ").strip()
        if not filepath:
            break
        if Path(filepath).exists():
            files.append(filepath)
        else:
            print(f"    ⚠️  File not found: {filepath}")

    if not files:
        print("❌ No valid files provided")
        sys.exit(1)

    config = {
        'name': name,
        'description': description or f"Documentation for {name}",
        'files': files
    }

    return config


def main():
    parser = argparse.ArgumentParser(
        description='Convert Word documents to Claude skills',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
Examples:
  # Interactive mode
  python3 cli/docx_scraper.py --interactive

  # Single document
  python3 cli/docx_scraper.py document.docx --name myskill --description "My documentation"

  # Multiple documents
  python3 cli/docx_scraper.py doc1.docx doc2.docx doc3.docx --name myskill

  # From config
  python3 cli/docx_scraper.py --config configs/doc_config.json

  # Skip extraction (rebuild from cached data)
  python3 cli/docx_scraper.py --name myskill --skip-extract
        '''
    )

    parser.add_argument(
        'files',
        nargs='*',
        help='Input .docx file path(s)'
    )

    parser.add_argument(
        '-n', '--name',
        help='Skill name'
    )

    parser.add_argument(
        '-d', '--description',
        help='Skill description'
    )

    parser.add_argument(
        '-c', '--config',
        help='Config JSON file'
    )

    parser.add_argument(
        '-i', '--interactive',
        action='store_true',
        help='Interactive mode'
    )

    parser.add_argument(
        '--skip-extract',
        action='store_true',
        help='Skip extraction, rebuild from cached data'
    )

    args = parser.parse_args()

    try:
        # Determine configuration source
        if args.interactive:
            config = interactive_mode()

        elif args.config:
            config = load_config_file(args.config)

        elif args.files and args.name:
            config = {
                'name': args.name,
                'description': args.description or f"Documentation for {args.name}",
                'files': args.files
            }

        elif args.skip_extract and args.name:
            # Rebuild from existing data
            config = {
                'name': args.name,
                'description': '',
                'files': []
            }

        else:
            parser.print_help()
            print("\nError: Provide either:")
            print("  1. --interactive")
            print("  2. --config <file>")
            print("  3. <files> --name <name>")
            print("  4. --name <name> --skip-extract")
            sys.exit(1)

        # Create converter
        converter = DocxToSkillConverter(config)

        # Process files
        if not args.skip_extract:
            if not config.get('files'):
                print("❌ No files to process")
                sys.exit(1)
            converter.process_all_files()

        # Build skill
        if args.skip_extract:
            # Load pages from cached data
            data_dir = Path(f"output/{config['name']}_data")
            if not data_dir.exists():
                print(f"❌ No cached data found in {data_dir}")
                sys.exit(1)

            pages_dir = data_dir / "pages"
            for page_file in pages_dir.glob("*.json"):
                with open(page_file, 'r', encoding='utf-8') as f:
                    converter.pages.append(json.load(f))

            print(f"✅ Loaded {len(converter.pages)} cached page(s)")

        converter.build_skill()

    except ImportError as e:
        print(f"❌ Error: {e}")
        print("\nInstall required library:")
        print("  pip install python-docx")
        sys.exit(1)

    except KeyboardInterrupt:
        print("\n\n⚠️  Interrupted by user")
        sys.exit(1)

    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
