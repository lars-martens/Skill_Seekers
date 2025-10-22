#!/usr/bin/env python3
"""
Simple .docx Text Extractor (Proof of Concept)

Task: B2.2 - Create simple .docx text extractor
Purpose: Demonstrate basic text extraction from Word documents
Usage: python3 cli/extract_docx_text.py <input.docx> [--output output.txt]

Features:
- Extract all paragraph text from .docx files
- Preserve paragraph structure
- Basic error handling
- Optional output to file

Requirements:
    pip install python-docx
"""

import sys
import argparse
from pathlib import Path


def extract_text_from_docx(filepath):
    """
    Extract all text from a .docx file.

    Args:
        filepath: Path to the .docx file

    Returns:
        str: Extracted text with paragraphs separated by double newlines

    Raises:
        ImportError: If python-docx is not installed
        FileNotFoundError: If the file doesn't exist
        Exception: For other docx parsing errors
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx library not found. Install with: pip install python-docx"
        )

    if not Path(filepath).exists():
        raise FileNotFoundError(f"File not found: {filepath}")

    try:
        doc = Document(filepath)
    except Exception as e:
        raise Exception(f"Failed to open .docx file: {e}")

    # Extract all paragraph text
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:  # Skip empty paragraphs
            paragraphs.append(text)

    return '\n\n'.join(paragraphs)


def get_document_stats(filepath):
    """
    Get basic statistics about the document.

    Args:
        filepath: Path to the .docx file

    Returns:
        dict: Statistics including paragraph count, character count, etc.
    """
    try:
        from docx import Document
    except ImportError:
        return {}

    doc = Document(filepath)

    # Count non-empty paragraphs
    paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    total_chars = sum(len(p.text) for p in paragraphs)
    total_words = sum(len(p.text.split()) for p in paragraphs)

    return {
        'paragraphs': len(paragraphs),
        'characters': total_chars,
        'words': total_words,
        'tables': len(doc.tables)
    }


def main():
    parser = argparse.ArgumentParser(
        description='Extract text from Microsoft Word (.docx) documents',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
Examples:
  # Extract and print to console
  python3 cli/extract_docx_text.py document.docx

  # Extract and save to file
  python3 cli/extract_docx_text.py document.docx --output output.txt

  # Show statistics
  python3 cli/extract_docx_text.py document.docx --stats
        '''
    )

    parser.add_argument(
        'input',
        help='Input .docx file path'
    )

    parser.add_argument(
        '-o', '--output',
        help='Output text file path (optional, prints to console if not specified)'
    )

    parser.add_argument(
        '-s', '--stats',
        action='store_true',
        help='Show document statistics'
    )

    parser.add_argument(
        '-v', '--verbose',
        action='store_true',
        help='Verbose output'
    )

    args = parser.parse_args()

    # Validate input file
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: File not found: {args.input}", file=sys.stderr)
        sys.exit(1)

    if input_path.suffix.lower() != '.docx':
        print(f"Warning: File doesn't have .docx extension: {args.input}", file=sys.stderr)

    try:
        # Extract text
        if args.verbose:
            print(f"Extracting text from: {args.input}", file=sys.stderr)

        text = extract_text_from_docx(args.input)

        # Show statistics if requested
        if args.stats:
            stats = get_document_stats(args.input)
            print(f"\nDocument Statistics:", file=sys.stderr)
            print(f"  Paragraphs: {stats.get('paragraphs', 0)}", file=sys.stderr)
            print(f"  Words: {stats.get('words', 0)}", file=sys.stderr)
            print(f"  Characters: {stats.get('characters', 0)}", file=sys.stderr)
            print(f"  Tables: {stats.get('tables', 0)}", file=sys.stderr)
            print(f"\nExtracted Text:", file=sys.stderr)
            print("-" * 40, file=sys.stderr)

        # Output text
        if args.output:
            output_path = Path(args.output)
            output_path.write_text(text, encoding='utf-8')
            if args.verbose:
                print(f"\nText saved to: {args.output}", file=sys.stderr)
        else:
            # Print to console
            print(text)

        if args.verbose:
            char_count = len(text)
            print(f"\nExtracted {char_count} characters", file=sys.stderr)

    except ImportError as e:
        print(f"Error: {e}", file=sys.stderr)
        print("\nInstall required library:", file=sys.stderr)
        print("  pip install python-docx", file=sys.stderr)
        sys.exit(1)

    except FileNotFoundError as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)

    except Exception as e:
        print(f"Error: Failed to extract text from .docx file", file=sys.stderr)
        print(f"Details: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == '__main__':
    main()
